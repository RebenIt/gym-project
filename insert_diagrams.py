#!/usr/bin/env python3
"""Insert the architecture and flowchart diagrams into the Word document."""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches

doc = Document('FitZone_Research_Project_Raparin_Institute_v2.docx')

# Find and replace figure placeholders with actual images
diagrams = {
    'System Architecture Diagram': 'assets/images/system_architecture.png',
    'Website Flowchart': 'assets/images/website_flowchart.png',
}

for ti, table in enumerate(doc.tables):
    cell_text = table.cell(0, 0).text if table.rows else ''
    for caption, img_path in diagrams.items():
        if caption in cell_text:
            # Clear the placeholder cell
            cell = table.cell(0, 0)
            for p in cell.paragraphs:
                p.clear()
            # Add image
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(5.5))
            print(f"Inserted: {caption}")

doc.save('FitZone_Research_Project_Raparin_Institute_v2.docx')
print("Document updated!")
