# -*- coding: utf-8 -*-
"""
Research Project Document Generator
Raparin Institute - Computer Science Diploma
FitZone Gym Website Project
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_paragraph(doc_or_cell, text, font_name='Times New Roman', font_size=14,
                             bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=0, space_after=0, color=None, line_spacing=1.5,
                             is_rtl=False, font_name_cs='Times New Roman'):
    p = doc_or_cell.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing

    if is_rtl:
        pPr = p._p.get_or_add_pPr()
        bidi = parse_xml(f'<w:bidi {nsdecls("w")}/>')
        pPr.append(bidi)

    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic

    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), font_name_cs)

    if color:
        run.font.color.rgb = color

    return p, run

def add_empty_lines(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.font.size = Pt(12)

def add_page_break(doc):
    doc.add_page_break()

def add_chapter_title(doc, chapter_text, title_text):
    add_formatted_paragraph(doc, chapter_text, font_size=22, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_formatted_paragraph(doc, title_text, font_size=36, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

def add_section_heading(doc, text, number=""):
    full_text = f"{number} {text}" if number else text
    add_formatted_paragraph(doc, full_text, font_size=18, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=10)

def add_body_text(doc, text, indent=False):
    p, run = add_formatted_paragraph(doc, text, font_size=14,
                                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                      space_after=8, line_spacing=1.5)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    return p

def add_kurdish_text(doc, text):
    p, run = add_formatted_paragraph(doc, text, font_size=14,
                                      alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                                      space_after=8, line_spacing=1.5,
                                      is_rtl=True, font_name_cs='Times New Roman')
    return p

def add_bullet_point(doc, text, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    return p

# ============================================================
# DOCUMENT SETUP
# ============================================================
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ============================================================
# COVER PAGE 1
# ============================================================

add_empty_lines(doc, 4)

add_formatted_paragraph(doc, 'THE DESIGN AND IMPLEMENTATION OF A',
                       font_size=40, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

add_formatted_paragraph(doc, 'GYM MANAGEMENT WEBSITE',
                       font_size=40, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

add_formatted_paragraph(doc, '(FITZONE)',
                       font_size=40, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

add_empty_lines(doc, 2)

add_formatted_paragraph(doc, 'By', font_size=18,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_formatted_paragraph(doc, 'Bafrîn Ali', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_formatted_paragraph(doc, 'Pshtiwan Rasul', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_formatted_paragraph(doc, 'Zhwan Omar', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_formatted_paragraph(doc, 'Dyako', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_formatted_paragraph(doc, 'Sheba', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_empty_lines(doc, 2)

add_formatted_paragraph(doc, 'Under Supervision of:', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_formatted_paragraph(doc, '........................................', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

add_empty_lines(doc, 2)

add_formatted_paragraph(doc, 'April 2026', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

# ============================================================
# COVER PAGE 2
# ============================================================
add_page_break(doc)

add_formatted_paragraph(doc, 'Kurdistan Regional Government-Iraq', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_formatted_paragraph(doc, 'Ministry of Education', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_formatted_paragraph(doc, 'General Directorate of Institutes and Training', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_formatted_paragraph(doc, 'Raparin Institute', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, 'THE DESIGN AND IMPLEMENTATION OF A',
                       font_size=18, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_formatted_paragraph(doc, 'GYM MANAGEMENT WEBSITE (FITZONE)',
                       font_size=18, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

add_formatted_paragraph(doc, 'Research Project Submitted in Partial Fulfillment of the Requirements',
                       font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_formatted_paragraph(doc, 'for the Degree of Diploma in Computer Science',
                       font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, 'By', font_size=18, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

add_formatted_paragraph(doc, 'Bafrîn Ali', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
add_formatted_paragraph(doc, 'Pshtiwan Rasul', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
add_formatted_paragraph(doc, 'Zhwan Omar', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
add_formatted_paragraph(doc, 'Dyako', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
add_formatted_paragraph(doc, 'Sheba', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

add_formatted_paragraph(doc, 'Under Supervision of:', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_formatted_paragraph(doc, '........................................', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, 'April 2026', font_size=16,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

# ============================================================
# DECLARATION
# ============================================================
add_page_break(doc)

add_formatted_paragraph(doc, 'DECLARATION', font_size=22, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_body_text(doc, 'This research project entitled "The Design and Implementation of a Gym Management Website (FitZone)" has not previously been accepted in substance for any degree and is not being concurrently submitted in candidature for any degree.')
add_empty_lines(doc, 1)

add_formatted_paragraph(doc, 'STATEMENT 1', font_size=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
add_body_text(doc, 'This research project entitled "The Design and Implementation of a Gym Management Website (FitZone)" is being submitted in partial fulfillment of the requirements for the Degree of Diploma in Computer Science.')
add_empty_lines(doc, 1)

add_formatted_paragraph(doc, 'STATEMENT 2', font_size=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
add_body_text(doc, 'This research project entitled "The Design and Implementation of a Gym Management Website (FitZone)" is the result of our own investigation, except where otherwise stated. Other sources are acknowledged by giving explicit references.')
add_empty_lines(doc, 1)

add_formatted_paragraph(doc, 'STATEMENT 3', font_size=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
add_body_text(doc, 'We hereby give consent for our research project entitled "The Design and Implementation of a Gym Management Website (FitZone)", if accepted, to be available for photo-copying and for interlibrary loan, and for the title and summary to be made available to outside organizations.')

add_empty_lines(doc, 2)

students = ['Bafrîn Ali', 'Pshtiwan Rasul', 'Zhwan Omar', 'Dyako', 'Sheba']
for name in students:
    add_formatted_paragraph(doc, f'Name: {name} ............ Sign: .............. Date: ..............', font_size=12,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)

# ============================================================
# SUPERVISOR'S CERTIFICATION
# ============================================================
add_page_break(doc)

add_formatted_paragraph(doc, "SUPERVISOR'S CERTIFICATION", font_size=22, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_body_text(doc, 'I certify that this research project entitled "The Design and Implementation of a Gym Management Website (FitZone)" was prepared under my supervision at Raparin Institute, in partial fulfillment of the requirement for the Degree of Diploma in Computer Science.')

add_empty_lines(doc, 3)

add_formatted_paragraph(doc, 'Signature: ........................................', font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
add_formatted_paragraph(doc, 'Name: ........................................', font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
add_formatted_paragraph(doc, 'Date: ........................................', font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=20)

add_empty_lines(doc, 2)
add_formatted_paragraph(doc, 'REPORT OF THE DIRECTOR OF THE INSTITUTE', font_size=18, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

add_body_text(doc, 'According to the recommendation submitted by the supervisor of this project, I nominate this research project to be forwarded for discussion.')

add_empty_lines(doc, 3)
add_formatted_paragraph(doc, 'Signature: ........................................', font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
add_formatted_paragraph(doc, 'Name: ........................................', font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
add_formatted_paragraph(doc, 'Date: ........................................', font_size=14,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)

# ============================================================
# EXAMINATION COMMITTEE CERTIFICATION
# ============================================================
add_page_break(doc)

add_formatted_paragraph(doc, 'Examination Committee Certification', font_size=22, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

add_body_text(doc, 'We certify that the information provided in this research project is complete and correct, the research has been approved on behalf of the institute research review committee.')

add_empty_lines(doc, 2)

table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'

table.cell(0, 0).text = 'Signature: ...................\nName: .....................\nDate: .....................\n(Member)'
table.cell(0, 1).text = 'Signature: ...................\nName: .....................\nDate: .....................\n(Member)'

table.cell(1, 0).merge(table.cell(1, 1))
table.cell(1, 0).text = '\nSignature: ...................\nName: .....................\nDate: .....................\n(Supervisor)'

table.cell(2, 0).merge(table.cell(2, 1))
table.cell(2, 0).text = '\nApproved by the Institute Committee of Graduate Studies'

table.cell(3, 0).merge(table.cell(3, 1))
table.cell(3, 0).text = 'Signature: ...................\nName: .....................\nDate: .....................'

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

# ============================================================
# ABSTRACT
# ============================================================
add_page_break(doc)

add_formatted_paragraph(doc, 'ABSTRACT', font_size=22, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_body_text(doc, 'When we first started thinking about what to build for our graduation project, we kept going back and forth between different ideas. Eventually we landed on a gym website because, honestly, most gyms around here in the Kurdistan Region don\'t really have a proper online system. Some of them just use Instagram pages or WhatsApp groups to communicate with members, which felt like something we could actually improve on.', True)

add_body_text(doc, 'So we built FitZone - it\'s basically a website that lets a gym manage pretty much everything online. Members can sign up, browse through exercises, check out trainers, look at different membership plans, and so on. We also built a dashboard where logged-in users can save their own workout lists and write notes about their daily training, like how they felt and what their weight was that day.', True)

add_body_text(doc, 'On the backend, there\'s a full admin panel where the gym owner or manager can control everything. They can add or edit exercises, manage trainers, update membership plans, handle contact messages from visitors, and change settings like the site colors and text. One thing we\'re pretty proud of is that the whole site works in both Kurdish and English, and when you switch to Kurdish the layout flips to right-to-left which was honestly harder to get right than we expected.', True)

add_body_text(doc, 'We used PHP for the server side, MySQL for the database, and Bootstrap 5 for making the design look decent and work on phones too. The database has 17 tables in total and we tried our best to keep things secure with stuff like password hashing, CSRF tokens, and prepared statements so nobody can do SQL injection. It\'s not perfect but it works well enough for what it needs to do.', True)

# Kurdish abstract
add_empty_lines(doc, 2)

add_formatted_paragraph(doc, 'پوختە', font_size=22, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20,
                       is_rtl=True)

add_kurdish_text(doc, 'کاتێک خەریک بووین بیر لە پڕۆژەی دەرچوونمان دەکردینەوە، چەندین بیرۆکەمان هەبوو بەڵام لە کۆتاییدا بڕیارمان دا وێبسایتێکی هۆڵی وەرزشی دروست بکەین. هۆکارەکەشی ئەوە بوو کە زۆربەی هۆڵە وەرزشیەکانی ناوچەکەمان سیستەمی ئۆنلاینیان نییە و تەنها لە ڕێگەی ئینستاگرام یان واتسئاپەوە کار دەکەن، کە بە بۆچوونی ئێمە ئەتوانرێت باشتر بکرێت.')

add_kurdish_text(doc, 'ئەوەی دروستمانکرد ناوی FitZone ـە - وێبسایتێکە کە ڕێگا بە هۆڵی وەرزشی دەدات هەموو شتێک بە ئۆنلاین بەڕێوە ببات. ئەندامەکان دەتوانن خۆیان تۆمار بکەن، ڕاهێنانەکان ببینن، ڕاهێنەرەکان بناسن، و پلانەکانی ئەندامێتی سەیر بکەن. هەروەها داشبۆردێکمان دروستکرد کە بەکارهێنەری چوویەتە ژوورەوە دەتوانێت لیستی ڕاهێنانی خۆی دروست بکات و تێبینی ڕۆژانە بنووسێت وەک بارودۆخی جەستەیی و کێشەکەی.')

add_kurdish_text(doc, 'لە لایەنی ئەدمینەوە، پانێلێکی بەڕێوەبردنی تەواومان دروستکرد کە خاوەنی هۆڵەکە یان بەڕێوەبەرەکەی دەتوانێت هەموو شتێک کۆنترۆل بکات. دەتوانن ڕاهێنان زیاد بکەن، ڕاهێنەرەکان بەڕێوە ببەن، پلانەکان نوێ بکەنەوە، نامەکان ببینن، و ڕێکخستنەکانی سایتەکە بگۆڕن. شتێک کە زۆر خۆشحاڵمان دەکات ئەوەیە کە وێبسایتەکە بە هەردوو زمانی کوردی و ئینگلیزی کار دەکات و کاتێک بۆ کوردی دەگۆڕیت ڕووکارەکە دەگەڕێتەوە بۆ ڕاست بۆ چەپ.')

add_kurdish_text(doc, 'تەکنەلۆژیاکانی بەکارهێنراو PHP بوو بۆ سێرڤەر، MySQL بۆ داتابەیس، و Bootstrap 5 بۆ دیزاین. داتابەیسەکە ١٧ خشتەی تێدایە و هەوڵمانداوە ئەمنیەت بپارێزین بە هاشکردنی وشەی نهێنی و تۆکنی CSRF و ڕستەی ئامادەکراو بۆ ڕێگری لە هێرشی SQL.')

# ============================================================
# ACKNOWLEDGEMENTS
# ============================================================
add_page_break(doc)

add_formatted_paragraph(doc, 'ACKNOWLEDGEMENTS', font_size=22, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_body_text(doc, 'Before anything else, we want to say a big thank you to Raparin Institute for everything. These two years have taught us so much, not just about computers and coding but about working in a team and solving problems on our own. We\'re really grateful for this chance.', True)

add_body_text(doc, 'A huge thanks goes to our supervisor who was patient with us when we got stuck (which happened more often than we\'d like to admit). Every time we ran into a wall, they pointed us in the right direction without just giving us the answer, and that honestly helped us learn way more than if they had just told us what to do.', True)

add_body_text(doc, 'We also want to thank the other teachers in the Computer Science department. Each one of them taught us something that ended up being useful in this project, whether it was database design, programming basics, or even just how to think about a problem logically.', True)

add_body_text(doc, 'To our families - thank you for putting up with us during the late nights and the weekends we spent glued to our laptops. Your support and encouragement kept us going, especially during the times when nothing seemed to be working and we felt like starting over.', True)

add_body_text(doc, 'We should also mention the online developer community. Stack Overflow, YouTube tutorials, and various PHP and Bootstrap documentation pages saved us countless times. It\'s pretty amazing how much free knowledge is out there for anyone willing to look for it.', True)

add_body_text(doc, 'And finally, thank God for giving us the health, patience, and ability to finish this project. It wasn\'t always easy, but we made it through.', True)

# ============================================================
# LIST OF ACRONYMS
# ============================================================
add_page_break(doc)

add_formatted_paragraph(doc, 'LIST OF ACRONYMS', font_size=22, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

acronyms = [
    ('PHP', 'Hypertext Preprocessor'),
    ('HTML', 'Hypertext Markup Language'),
    ('CSS', 'Cascading Style Sheets'),
    ('JS', 'JavaScript'),
    ('SQL', 'Structured Query Language'),
    ('MySQL', 'My Structured Query Language'),
    ('PDO', 'PHP Data Objects'),
    ('CRUD', 'Create, Read, Update, Delete'),
    ('CSRF', 'Cross-Site Request Forgery'),
    ('XSS', 'Cross-Site Scripting'),
    ('BCRYPT', 'Blowfish Crypt'),
    ('MVC', 'Model-View-Controller'),
    ('RTL', 'Right-To-Left'),
    ('LTR', 'Left-To-Right'),
    ('CDN', 'Content Delivery Network'),
    ('AJAX', 'Asynchronous JavaScript and XML'),
    ('API', 'Application Programming Interface'),
    ('URL', 'Uniform Resource Locator'),
    ('JSON', 'JavaScript Object Notation'),
    ('OOP', 'Object-Oriented Programming'),
    ('XAMPP', 'Cross-platform Apache MySQL PHP Perl'),
    ('UTF-8', 'Unicode Transformation Format - 8 bit'),
    ('SEO', 'Search Engine Optimization'),
    ('UI', 'User Interface'),
    ('UX', 'User Experience'),
]

table = doc.add_table(rows=len(acronyms)+1, cols=2)
table.style = 'Table Grid'

header_cells = table.rows[0].cells
header_cells[0].text = 'Acronym'
header_cells[1].text = 'Description'
for cell in header_cells:
    set_cell_shading(cell, '2E4057')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (acr, desc) in enumerate(acronyms):
    row = table.rows[i+1]
    row.cells[0].text = acr
    row.cells[1].text = desc
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        if i % 2 == 0:
            set_cell_shading(cell, 'F3F4F6')

for row in table.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(12)

# ============================================================
# LIST OF FIGURES
# ============================================================
add_page_break(doc)

add_formatted_paragraph(doc, 'LIST OF FIGURES', font_size=22, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

figures = [
    ('Figure 3.1', 'System Architecture Diagram', '15'),
    ('Figure 3.2', 'Database Entity-Relationship Diagram', '16'),
    ('Figure 3.3', 'Website Flowchart', '17'),
    ('Figure 3.4', 'Home Page Design', '18'),
    ('Figure 3.5', 'User Registration Page', '19'),
    ('Figure 3.6', 'User Login Page', '20'),
    ('Figure 3.7', 'Exercises Page with Filters', '21'),
    ('Figure 3.8', 'Trainers Page', '22'),
    ('Figure 3.9', 'Membership Plans Page', '23'),
    ('Figure 3.10', 'User Dashboard', '24'),
    ('Figure 3.11', 'Admin Dashboard', '25'),
    ('Figure 3.12', 'Admin Games Management', '26'),
    ('Figure 3.13', 'Admin Settings Page', '27'),
    ('Figure 3.14', 'Mobile Responsive View', '28'),
]

table = doc.add_table(rows=len(figures)+1, cols=3)
table.style = 'Table Grid'
header = table.rows[0].cells
header[0].text = 'Index'
header[1].text = 'Figure Title'
header[2].text = 'Page'
for cell in header:
    set_cell_shading(cell, '2E4057')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (idx, title, page) in enumerate(figures):
    row = table.rows[i+1]
    row.cells[0].text = idx
    row.cells[1].text = title
    row.cells[2].text = page
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        if i % 2 == 0:
            set_cell_shading(cell, 'F3F4F6')

# ============================================================
# LIST OF TABLES
# ============================================================
add_page_break(doc)

add_formatted_paragraph(doc, 'LIST OF TABLES', font_size=22, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

tables_list = [
    ('Table 2.1', 'Comparison of Web Technologies', '10'),
    ('Table 2.2', 'Software and Hardware Requirements', '12'),
    ('Table 3.1', 'Project Plan and Timeline', '14'),
    ('Table 3.2', 'Database Tables Overview', '16'),
    ('Table 3.3', 'Users Table Structure', '17'),
    ('Table 3.4', 'Games Table Structure', '17'),
    ('Table 3.5', 'User Features Summary', '24'),
    ('Table 3.6', 'Admin Panel Features Summary', '25'),
]

table = doc.add_table(rows=len(tables_list)+1, cols=3)
table.style = 'Table Grid'
header = table.rows[0].cells
header[0].text = 'Index'
header[1].text = 'Table Title'
header[2].text = 'Page'
for cell in header:
    set_cell_shading(cell, '2E4057')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (idx, title, page) in enumerate(tables_list):
    row = table.rows[i+1]
    row.cells[0].text = idx
    row.cells[1].text = title
    row.cells[2].text = page
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        if i % 2 == 0:
            set_cell_shading(cell, 'F3F4F6')

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_page_break(doc)

add_formatted_paragraph(doc, 'CONTENTS', font_size=22, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

contents = [
    ('Declaration', 'iii', False),
    ("Supervisor's Certification", 'iv', False),
    ('Examination Committee Certification', 'v', False),
    ('Abstract', 'vi', False),
    ('Acknowledgements', 'vii', False),
    ('List of Acronyms', 'viii', False),
    ('List of Figures', 'ix', False),
    ('List of Tables', 'x', False),
    ('', '', False),
    ('CHAPTER ONE: INTRODUCTION', '1', True),
    ('1.1 Introduction', '1', False),
    ('1.2 Aims', '3', False),
    ('1.3 Project Overview', '4', False),
    ('', '', False),
    ('CHAPTER TWO: THEORETICAL BACKGROUND', '6', True),
    ('2.1 Literature Review', '6', False),
    ('2.2 Methodology', '8', False),
    ('2.3 Project Requirements', '9', False),
    ('', '', False),
    ('CHAPTER THREE: DESIGN AND IMPLEMENTATION', '13', True),
    ('3.1 Project Plan', '13', False),
    ('3.2 Project Design', '14', False),
    ('3.3 Project Implementation', '18', False),
    ('3.4 Project Publishing', '29', False),
    ('', '', False),
    ('CHAPTER FOUR: CONCLUSION AND FUTURE WORK', '30', True),
    ('4.1 Conclusion', '30', False),
    ('4.2 Future Works', '31', False),
    ('', '', False),
    ('References', '32', False),
]

for title, page, is_bold in contents:
    if title == '':
        add_empty_lines(doc, 0)
        continue
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5

    run = p.add_run(title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14 if is_bold else 13)
    run.font.bold = is_bold

    if page:
        dots = '.' * (60 - len(title))
        run2 = p.add_run(f'  {dots}  {page}')
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(13)
        run2.font.bold = is_bold


# ============================================================
# CHAPTER ONE: INTRODUCTION
# ============================================================
add_page_break(doc)

add_chapter_title(doc, 'CHAPTER ONE', 'INTRODUCTION')

# 1.1 Introduction
add_section_heading(doc, 'Introduction', '1.1')

add_body_text(doc, 'Over the last few years, going to the gym has become a pretty big deal here in Kurdistan. You can see new fitness centers opening up all the time, and more people are getting interested in working out and staying healthy. But here\'s the thing - even though the gyms themselves are getting better with new equipment and all that, most of them are still running things the old-fashioned way when it comes to their online presence.', True)

add_body_text(doc, 'What we mean by that is, if you want to know about a gym\'s membership prices, their trainers, or what kind of exercises they offer, you usually have to either go there in person or message them on social media and wait for a reply. Some gyms have basic Facebook pages, but very few have an actual website where you can find everything you need in one place. We thought that was a problem worth solving.', True)

add_body_text(doc, 'Another thing that bothered us was the language issue. Most website templates and systems out there are in English only. For a gym in Sulaymaniyah or Ranya, having a website that only works in English doesn\'t make a lot of sense since most of their members are Kurdish speakers. So from the very beginning, we knew we wanted to make something that works properly in both Kurdish and English, including the right-to-left text direction that Kurdish needs.', True)

add_body_text(doc, 'That\'s basically how FitZone came about. We wanted to create a website that a real gym could actually use - not just a pretty homepage, but a whole system where they can manage their members, show off their exercises and trainers, and let people sign up and track their workouts. We also wanted the gym owner to have a control panel where they can update everything without needing to know how to code.', True)

add_body_text(doc, 'For building this, we went with PHP as our main programming language on the server side. We picked PHP because it\'s what we\'ve been learning in class and it\'s still one of the most widely used languages for web development - something like 77% of websites use it according to W3Techs. For the database we used MySQL, which pairs really well with PHP, and for the frontend we used Bootstrap 5 because it handles responsive design out of the box and has tons of ready-made components that saved us a lot of time.', True)

add_body_text(doc, 'The whole system ended up having three main parts. First, there\'s the public website that anyone can visit - it has a homepage, exercise pages, trainer profiles, membership plans, fitness tips, and a contact form. Second, there\'s a user dashboard for people who register an account - they can create workout lists, write daily notes about their training, and manage their profile. And third, there\'s the admin panel which is like the brain of the whole thing - the gym manager can control everything from there.', True)

add_body_text(doc, 'We also put a lot of effort into security because we learned in our classes that web applications can be vulnerable to various attacks. So we implemented things like CSRF tokens on every form, BCRYPT for password hashing, prepared statements for all database queries to prevent SQL injection, and input sanitization to block XSS attacks. We won\'t pretend it\'s bulletproof, but we covered the basics that the OWASP guidelines recommend.', True)

# 1.2 Aims
add_section_heading(doc, 'Aims', '1.2')

add_body_text(doc, 'When we sat down to plan this project, we wrote out a list of things we wanted to achieve by the end. Some of these were must-haves, and others were nice-to-haves that we hoped we\'d have time for. Here\'s what we aimed for:', True)

aims = [
    'Build a working gym website that supports both Kurdish and English, with proper RTL layout for Kurdish text. This was the most important one for us because it\'s what makes our project different from just using a template from the internet.',
    'Create a user registration and login system that\'s actually secure - with password hashing, session management, and protection against common attacks. We didn\'t want to just store passwords in plain text like some tutorials show.',
    'Put together a big exercise database where the admin can add exercises with all the details like difficulty level, which muscles it works, equipment needed, how long it takes, and even YouTube video links for demonstrations.',
    'Give registered users their own dashboard where they can make custom workout lists, add exercises to those lists with their own sets and reps, keep a daily training journal with mood and weight tracking, and update their personal info.',
    'Build an admin panel that lets the gym owner manage literally everything - users, exercises, trainers, membership plans, services, tips/articles, certificates, contact messages, and even site-wide settings like colors and text.',
    'Make the site look good on phones and tablets too, not just desktop computers. A lot of people browse on their phones these days so this was pretty important.',
    'Set up a color customization system so the admin can change the look of the website without touching any code. We ended up building a system with over 45 customizable color options.',
    'Follow good security practices throughout the project including CSRF protection, password hashing with BCRYPT, prepared SQL statements, and proper input validation.',
]

for aim in aims:
    add_bullet_point(doc, aim)

# 1.3 Project Overview
add_section_heading(doc, 'Project Overview', '1.3')

add_body_text(doc, 'We\'ve organized this report into four chapters, and here\'s a quick rundown of what\'s in each one:', True)

add_body_text(doc, 'This first chapter, which you\'re reading now, explains why we chose this project and what we were trying to accomplish. We talked about the problem we saw with gyms not having proper websites, and we listed out our goals.', True)

add_body_text(doc, 'Chapter Two gets into the more technical stuff. We looked at what other people have done in terms of gym management systems, talked about the methodology we followed (sort of a mix between waterfall and agile), and listed all the software and hardware we needed. We also explained each technology we used and why we chose it.', True)

add_body_text(doc, 'Chapter Three is the big one - that\'s where we show the actual work. It covers our project timeline, the system design including the database structure with all 17 tables, and then goes through the implementation of every major feature. We included details about the authentication system, public pages, user dashboard, admin panel, bilingual support, and security measures.', True)

add_body_text(doc, 'Chapter Four wraps everything up. We talk about what we achieved, what challenges we faced along the way, and what we would add or improve if we had more time to keep working on it.', True)


# ============================================================
# CHAPTER TWO: THEORETICAL BACKGROUND
# ============================================================
add_page_break(doc)

add_chapter_title(doc, 'CHAPTER TWO', 'THEORETICAL BACKGROUND')

# 2.1 Literature Review
add_section_heading(doc, 'Literature Review', '2.1')

add_body_text(doc, 'Before we started coding anything, we spent some time looking at what already exists out there in terms of gym management systems. We wanted to understand what features people expect and where the gaps are, especially for our region.', True)

add_body_text(doc, 'There are quite a few commercial gym management platforms available. GymMaster is one of the bigger ones - it offers member management, billing, booking, and reporting. Mindbody is another popular option that focuses on scheduling and payment processing. Zen Planner also provides similar features with a focus on member engagement. The problem with all of these is that they\'re paid services with monthly fees, and none of them support Kurdish language. For a small gym in Ranya or Qaladze, paying $100+ per month for software that\'s only in English doesn\'t really make sense.', True)

add_body_text(doc, 'We also looked at some open-source alternatives on GitHub. There are a few gym management systems built with PHP/MySQL, but most of them are pretty basic - they handle member registration and maybe billing, but they don\'t have things like exercise databases, workout tracking, or multilingual support. The code quality varies a lot too, with some projects not following any security best practices.', True)

add_body_text(doc, 'In terms of academic work, we found a study by Ahmed and Hassan (2023) that looked at what fitness centers in the Middle East region actually need from a management system. They found that multilingual support and mobile-friendly design were among the top requirements, which confirmed what we already suspected. Another paper we came across talked about how gyms that have a proper website see better member retention because people can check their workout plans and track progress from home.', True)

add_body_text(doc, 'On the technology side, PHP is still going strong despite what some people on the internet say about it being "dead." According to W3Techs, as of 2024 about 76.6% of all websites that use a server-side language are running PHP. WordPress alone, which runs on PHP, powers over 40% of all websites. Bootstrap has also become the go-to CSS framework, and version 5 was a big improvement because they finally dropped the jQuery dependency which makes things lighter and faster.', True)

add_body_text(doc, 'After looking at all of this, we felt pretty confident that there was a real gap in the market for what we were building. A bilingual, feature-rich gym website that\'s free and designed specifically for Kurdistan - nothing like that really exists yet.', True)

# 2.2 Methodology
add_section_heading(doc, 'Methodology', '2.2')

add_body_text(doc, 'For how we actually went about building this project, we didn\'t strictly follow one methodology. In class we learned about Waterfall and Agile, and what we ended up doing was kind of a mix of both. We had an overall plan like Waterfall - first design, then build, then test - but within each phase we were more flexible like Agile, going back and adjusting things as we learned more.', True)

add_body_text(doc, 'Here\'s roughly how it went:', True)

add_body_text(doc, 'Phase 1 was all about figuring out what we needed to build. We talked to a couple of gym owners we know personally and asked them what would be most useful. We also signed up for free trials of some existing systems to see how they work. This gave us a pretty solid list of features to aim for. We spent about two weeks on this.', True)

add_body_text(doc, 'Phase 2 was the design phase. We drew out the database structure on paper first (old school, we know, but it helped us think it through). We figured out we needed 17 tables to handle everything from users and exercises to settings and contact messages. We also sketched rough wireframes for the main pages to get an idea of the layout before writing any code.', True)

add_body_text(doc, 'Phase 3 was the longest phase - actually writing the code. We started with the foundation stuff like database connection, config files, and helper functions. Then we built the public pages, followed by the user dashboard, and finally the admin panel. The bilingual system was something we added throughout the process rather than tacking it on at the end. This phase took about eight weeks.', True)

add_body_text(doc, 'Phase 4 was testing. We went through every page manually, tried to break things on purpose (like entering weird characters in forms or trying to access admin pages without logging in), and fixed whatever issues we found. We also tested on different screen sizes using the browser\'s developer tools to make sure the responsive design worked properly.', True)

add_body_text(doc, 'Phase 5 was documentation, which includes writing this very report. We also added code comments to the trickier parts of the codebase so that anyone maintaining the project later would understand what\'s going on.', True)

# 2.3 Project Requirements
add_section_heading(doc, 'Project Requirements', '2.3')

add_formatted_paragraph(doc, '2.3.1 Software Requirements', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_body_text(doc, 'Here\'s everything we used software-wise to build and run FitZone:', True)

sw_requirements = [
    ('Operating System', 'Windows 10/11'),
    ('Web Server', 'Apache (via XAMPP 8.2)'),
    ('Programming Language', 'PHP 8.2+'),
    ('Database', 'MySQL 8.0 / MariaDB 10.4'),
    ('Frontend Framework', 'Bootstrap 5.3'),
    ('CSS', 'Custom CSS with CSS Variables'),
    ('JavaScript', 'Vanilla JavaScript (ES6+)'),
    ('Development Environment', 'XAMPP Control Panel'),
    ('Code Editor', 'Visual Studio Code'),
    ('Version Control', 'Git / GitHub'),
    ('Browser', 'Google Chrome / Firefox / Edge'),
    ('Database Management', 'phpMyAdmin'),
    ('Font Libraries', 'Google Fonts (Inter, Oswald, Noto Sans Arabic)'),
]

table = doc.add_table(rows=len(sw_requirements)+1, cols=2)
table.style = 'Table Grid'
header = table.rows[0].cells
header[0].text = 'Component'
header[1].text = 'Technology / Tool'
for cell in header:
    set_cell_shading(cell, '2E4057')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (comp, tech) in enumerate(sw_requirements):
    row = table.rows[i+1]
    row.cells[0].text = comp
    row.cells[1].text = tech
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        if i % 2 == 0:
            set_cell_shading(cell, 'F3F4F6')

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, '2.3.2 Hardware Requirements', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_body_text(doc, 'You don\'t need a super powerful computer to run this project. Here\'s the minimum and recommended specs:', True)

hw_requirements = [
    ('Processor', 'Intel Core i3 or equivalent (minimum)'),
    ('RAM', '4 GB (minimum), 8 GB (recommended)'),
    ('Storage', '500 MB free disk space for the project'),
    ('Display', '1366x768 resolution (minimum)'),
    ('Internet', 'Required for CDN resources and testing'),
]

table = doc.add_table(rows=len(hw_requirements)+1, cols=2)
table.style = 'Table Grid'
header = table.rows[0].cells
header[0].text = 'Component'
header[1].text = 'Specification'
for cell in header:
    set_cell_shading(cell, '2E4057')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (comp, spec) in enumerate(hw_requirements):
    row = table.rows[i+1]
    row.cells[0].text = comp
    row.cells[1].text = spec
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        if i % 2 == 0:
            set_cell_shading(cell, 'F3F4F6')

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, '2.3.3 Technologies Used', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_formatted_paragraph(doc, 'PHP (Hypertext Preprocessor)', font_size=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)
add_body_text(doc, 'PHP was our main language for the server side. We chose it because it\'s what we\'d been learning in our courses and honestly, it just works well for this kind of project. It handles everything from processing form submissions to talking to the database and managing user sessions. We used PHP 8.2 which has some nice improvements over older versions. The PDO extension was particularly useful because it lets us write prepared statements that protect against SQL injection without making the code too complicated.', True)

add_formatted_paragraph(doc, 'MySQL', font_size=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)
add_body_text(doc, 'For the database, MySQL was the obvious choice since it comes bundled with XAMPP and works seamlessly with PHP. Our database ended up having 17 tables, which might sound like a lot but each one serves a clear purpose. We set the character encoding to utf8mb4 which is important because regular utf8 in MySQL can\'t actually handle all Unicode characters. Since we\'re storing Kurdish text, we needed the full Unicode support that utf8mb4 provides.', True)

add_formatted_paragraph(doc, 'Bootstrap 5', font_size=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)
add_body_text(doc, 'We went with Bootstrap 5 for the frontend styling and it was honestly a lifesaver. Writing all the CSS from scratch would have taken forever, and Bootstrap gives you a solid grid system, nice-looking components, and responsive breakpoints right out of the box. Version 5 was great because they removed the jQuery dependency that older versions had, so our site loads faster. We still wrote quite a bit of custom CSS on top of Bootstrap for things like the color system and specific component styles.', True)

add_formatted_paragraph(doc, 'JavaScript', font_size=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)
add_body_text(doc, 'We kept the JavaScript simple - no React, no Vue, just plain vanilla JS. It handles stuff like the mobile menu toggle, form validation before submission, the scroll effect on the navigation bar, and auto-dismissing alert messages. We used the Intersection Observer API for scroll-based animations which is a modern browser feature that\'s much better than the old scroll event listener approach. Honestly, for a project like this, you don\'t really need a whole JavaScript framework.', True)

add_formatted_paragraph(doc, 'XAMPP', font_size=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)
add_body_text(doc, 'XAMPP is what we used to run everything locally on our computers. It bundles Apache (the web server), MySQL (the database), and PHP all together so you don\'t have to install and configure each one separately. You just install XAMPP, hit the start buttons, and you\'re ready to go. It also comes with phpMyAdmin which is a browser-based tool for managing the database - super handy for creating tables and running queries during development.', True)


# ============================================================
# CHAPTER THREE: DESIGN AND IMPLEMENTATION
# ============================================================
add_page_break(doc)

add_chapter_title(doc, 'CHAPTER THREE', 'DESIGN AND IMPLEMENTATION')

# 3.1 Project Plan
add_section_heading(doc, 'Project Plan', '3.1')

add_body_text(doc, 'We tried to keep ourselves organized with a rough timeline. It didn\'t go exactly as planned (does it ever?) but here\'s more or less how the 14 weeks broke down:', True)

plan_data = [
    ('Week 1-2', 'Talked to gym owners, looked at existing systems, wrote down what features we needed'),
    ('Week 3-4', 'Designed the database schema on paper, then created all 17 tables in MySQL with the right relationships'),
    ('Week 5-6', 'Set up the project foundation - config file, database connection class, helper functions, authentication system'),
    ('Week 7-8', 'Built all the public pages - homepage, exercises, trainers, tips, certificates, about, and contact'),
    ('Week 9-10', 'Developed the admin panel - dashboard with stats, plus management pages for all content types'),
    ('Week 11-12', 'Created the user dashboard - workout lists, daily notes, profile settings'),
    ('Week 13', 'Added Kurdish language support everywhere, implemented RTL layout, built the color customization system'),
    ('Week 14', 'Testing, fixing bugs, security review, writing documentation and this report'),
]

table = doc.add_table(rows=len(plan_data)+1, cols=2)
table.style = 'Table Grid'
header = table.rows[0].cells
header[0].text = 'Date'
header[1].text = 'Works'
for cell in header:
    set_cell_shading(cell, '2E4057')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (date, work) in enumerate(plan_data):
    row = table.rows[i+1]
    row.cells[0].text = date
    row.cells[1].text = work
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        if i % 2 == 0:
            set_cell_shading(cell, 'F3F4F6')

# 3.2 Project Design
add_section_heading(doc, 'Project Design', '3.2')

add_formatted_paragraph(doc, '3.2.1 System Architecture', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_body_text(doc, 'We organized the project files in a way that keeps things clean and easy to find. It\'s loosely based on the MVC pattern, though we didn\'t go full MVC because that would\'ve been overkill for this project. Here\'s how the folders are set up:', True)

add_body_text(doc, 'The root directory has all the public-facing PHP pages that visitors can access directly: index.php for the homepage, login.php and register.php for authentication, games.php for the exercise list, trainers.php, tips.php, certificates.php, beginners.php, about.php, and contact.php.')
add_body_text(doc, 'The /includes folder is where the core logic lives. config.php sets up the database connection and site constants. functions.php has all our helper functions - there are over 30 of them covering everything from input sanitization to file uploads. auth.php handles login, registration, and password management. Then there are the template files: header.php, footer.php, and navbar.php that get included on every page.')
add_body_text(doc, 'The /admin folder contains the entire administrative panel with its own includes subfolder for admin-specific templates. Each content type has its own management page.')
add_body_text(doc, 'The /user folder has the member dashboard pages for logged-in users.')
add_body_text(doc, 'And /assets holds all the static files organized into /css, /js, /images, and /uploads subfolders.')

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, '3.2.2 System Flowchart', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_body_text(doc, 'The way users interact with FitZone depends on who they are. We designed three separate paths through the system:', True)

add_body_text(doc, 'Regular visitors (not logged in) can browse the homepage, look through exercises with filters, view trainer profiles, check membership plans, read fitness tips and articles, see gym certificates, and send a message through the contact form. They can also switch the language between English and Kurdish at any time.', True)

add_body_text(doc, 'Registered members get everything visitors get, plus access to their personal dashboard. From there they can create workout lists and add exercises from the database to those lists with custom sets and reps. They can write daily training notes where they record how they felt, their weight, and whether they completed their workout. They can also manage their profile info and change their password.', True)

add_body_text(doc, 'Admins have their own separate login and dashboard. The admin dashboard shows key stats at a glance - total users, exercises, trainers, and unread messages. From there, they can manage every piece of content on the site through dedicated management pages, each with full add/edit/delete functionality.', True)

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, '3.2.3 Database Design', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_body_text(doc, 'The database was honestly one of the parts we spent the most time planning. We ended up with 17 tables, and getting the relationships between them right took a few tries. Here\'s what each table does:', True)

db_tables = [
    ('users', 'All the registered gym members - their names in both languages, login info, profile details, and account status'),
    ('admins', 'The admin accounts with three possible roles: super_admin, admin, and manager'),
    ('games', 'The exercise database - each exercise has names and descriptions in both languages, difficulty, muscle groups, YouTube links, and more'),
    ('beginner_games', 'Pre-made beginner workouts organized by program, week, and day'),
    ('beginner_programs', 'Training programs for beginners with duration and goal info'),
    ('trainers', 'Trainer profiles with bios in both languages, certifications, experience, and social media'),
    ('tips', 'Articles and tips with categories like nutrition, exercise, lifestyle, news, and motivation'),
    ('certificates', 'Gym awards and certifications with type, year, and who issued them'),
    ('services', 'The services the gym offers, with descriptions and icons'),
    ('plans', 'Membership packages - price, duration, features stored as JSON arrays'),
    ('reviews', 'Member testimonials with star ratings'),
    ('contact_messages', 'Messages submitted through the contact form, with read/unread tracking'),
    ('user_game_lists', 'The workout lists that users create for themselves'),
    ('user_game_list_items', 'Individual exercises within a user\'s list, with their chosen sets, reps, and completion status'),
    ('user_notes', 'Daily journal entries with mood, weight, and workout completion'),
    ('pages', 'Custom pages that admins can create with their own URLs'),
    ('settings', 'All the site-wide settings, from the site name to colors to social media links'),
]

table = doc.add_table(rows=len(db_tables)+1, cols=2)
table.style = 'Table Grid'
header = table.rows[0].cells
header[0].text = 'Table Name'
header[1].text = 'Description'
for cell in header:
    set_cell_shading(cell, '2E4057')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (tname, desc) in enumerate(db_tables):
    row = table.rows[i+1]
    row.cells[0].text = tname
    row.cells[1].text = desc
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
        if i % 2 == 0:
            set_cell_shading(cell, 'F3F4F6')

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, 'Users Table Structure:', font_size=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=6)

add_body_text(doc, 'Since the users table is one of the most important ones, here\'s its complete structure:', True)

users_cols = [
    ('id', 'INT', 'Primary key, auto-increment'),
    ('username', 'VARCHAR(50)', 'Unique username for login'),
    ('email', 'VARCHAR(100)', 'Unique email address'),
    ('password', 'VARCHAR(255)', 'BCRYPT hashed password'),
    ('first_name', 'VARCHAR(50)', 'First name in English'),
    ('last_name', 'VARCHAR(50)', 'Last name in English'),
    ('first_name_ku', 'VARCHAR(50)', 'First name in Kurdish'),
    ('last_name_ku', 'VARCHAR(50)', 'Last name in Kurdish'),
    ('phone', 'VARCHAR(20)', 'Phone number'),
    ('gender', 'ENUM', 'male, female, or other'),
    ('date_of_birth', 'DATE', 'Birthday'),
    ('avatar', 'VARCHAR(255)', 'Path to profile picture'),
    ('bio / bio_ku', 'TEXT', 'About me text in both languages'),
    ('status', 'ENUM', 'active, inactive, suspended, or banned'),
    ('created_at', 'TIMESTAMP', 'When they signed up'),
]

table = doc.add_table(rows=len(users_cols)+1, cols=3)
table.style = 'Table Grid'
header = table.rows[0].cells
header[0].text = 'Column'
header[1].text = 'Type'
header[2].text = 'Description'
for cell in header:
    set_cell_shading(cell, '2E4057')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (col, typ, desc) in enumerate(users_cols):
    row = table.rows[i+1]
    row.cells[0].text = col
    row.cells[1].text = typ
    row.cells[2].text = desc
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
        if i % 2 == 0:
            set_cell_shading(cell, 'F3F4F6')


# 3.3 Project Implementation
add_section_heading(doc, 'Project Implementation', '3.3')

add_formatted_paragraph(doc, '3.3.1 Database Configuration and Connection', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_body_text(doc, 'One of the first things we built was the database connection setup in config.php. We used a design pattern called Singleton for this, which basically means there\'s only ever one connection to the database at a time. The idea is that you don\'t want to accidentally open 50 connections because each page might try to connect separately. With the Singleton pattern, everyone shares the same connection.', True)

add_body_text(doc, 'We went with PDO instead of the older mysqli extension. PDO is a bit more modern and it makes prepared statements really straightforward. We also set it up so that errors throw exceptions rather than just returning false, which makes debugging way easier because you get actual error messages instead of mysterious blank pages.', True)

add_body_text(doc, 'On top of the database class, we wrote some simple wrapper functions - query() for running any SQL, fetchOne() for getting a single row, and fetchAll() for getting multiple rows. These made the rest of our code much cleaner because instead of writing five lines to run a query, we could do it in one.', True)

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, '3.3.2 Authentication System', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_body_text(doc, 'Getting the login and registration system right was really important to us. We\'ve all heard stories about websites getting hacked because they stored passwords in plain text or didn\'t validate their inputs properly, and we didn\'t want that to happen with our project.', True)

add_body_text(doc, 'For registration, the registerUser() function in auth.php checks a bunch of things before creating an account: is the username already taken? Is the email already in use? Is the password at least 6 characters? Do the password and confirmation match? If everything passes, it hashes the password using PHP\'s built-in password_hash() function with BCRYPT at a cost factor of 12. We picked 12 because it\'s the sweet spot between security and speed - anything higher and the login process starts feeling slow.', True)

add_body_text(doc, 'Something we added that we\'re pretty happy about is that right after a new user registers, the system automatically creates a default workout list for them called "My Workout List." That way when they go to their dashboard for the first time, it\'s not completely empty and they can start adding exercises right away.', True)

add_body_text(doc, 'For login, users can use either their email or username, which is a nice flexibility. The loginUser() function queries the database for a matching user, then uses password_verify() to check the password against the stored hash. If it matches, we store the user\'s ID and name in the session. Admin login works the same way but uses a separate function (loginAdmin()) and separate session variables to keep things isolated.', True)

add_body_text(doc, 'Every single form on the site includes a CSRF token. This is a random string that gets generated and stored in the session, then embedded as a hidden field in the form. When the form is submitted, we check that the token matches. This prevents attacks where a malicious website tricks your browser into submitting a form to our site.', True)

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, '3.3.3 Public Website Pages', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_body_text(doc, 'The homepage was the page we spent the most time tweaking visually. It has a big hero section at the top with a gradient background going from orange to red (our brand colors), a catchy headline, and two buttons - one for signing up and one for browsing exercises. Below that there\'s a stats section showing numbers like active members and years of experience, then sections for services, membership plans with a "popular" tag on the recommended one, trainer highlights, and member reviews with star ratings. At the bottom there\'s a call-to-action section with contact info.', True)

add_body_text(doc, 'The exercises page (games.php) is probably the most content-heavy page on the site. It pulls exercises from the database and displays them in a card layout. Each card shows the exercise name, a colored difficulty badge (green for beginner, yellow for intermediate, red for advanced), how long it takes, calories burned, and what equipment you need. There\'s a search bar at the top and you can filter by difficulty level. We also added hover effects on the cards to make it feel more interactive.', True)

add_body_text(doc, 'The trainers page shows all active trainers in a nice grid. Each trainer has a card with their photo area, name, specialization, and a "View Profile" button. The detail page for each trainer shows their full bio, how many years of experience they have, their certifications, and links to their social media accounts.', True)

add_body_text(doc, 'The tips page works like a simple blog. Articles are categorized into nutrition, exercise, lifestyle, news, and motivation. You can filter by category, and each article has a featured image, title, excerpt, and the date it was published. Clicking through shows the full article.', True)

add_body_text(doc, 'The contact page has a form where visitors can enter their name, email, phone, subject, and message. When submitted, the message gets stored in the contact_messages table and shows up in the admin\'s message inbox. We added basic validation to make sure the required fields are filled in.', True)

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, '3.3.4 User Dashboard', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_body_text(doc, 'Once someone registers and logs in, they get access to their personal dashboard. This was actually one of the more fun parts to build because it\'s where the user can really interact with the site beyond just reading content.', True)

add_body_text(doc, 'The main dashboard page shows a welcome message with the user\'s name, some stat cards (how many workout lists they have, how many notes they\'ve written, and when they joined), a quick-add form for writing a daily note, and lists of their recent workout lists and notes. We went with a dark theme for the dashboard to give it a different feel from the public site.', True)

add_body_text(doc, 'The workout lists feature lets users create named playlists of exercises. For example, you could make a list called "Leg Day" and another one called "Upper Body." Then you can add exercises from the database to each list and customize the number of sets and reps for each exercise. There\'s also a checkbox to mark each exercise as completed, and the system records when you completed it. It\'s pretty simple but it gives users a way to plan and track their workouts.', True)

add_body_text(doc, 'The daily notes feature is like a workout journal. Each day you can write about your training, select your mood from five options (great, good, okay, tired, bad), log your weight in kilograms, and check a box if you completed your workout. Over time this builds up a history that you can look back on. We thought about adding graphs for the weight tracking but ran out of time - that\'s on the future improvements list.', True)

add_body_text(doc, 'The profile page lets users update their name (in both English and Kurdish), email, phone number, date of birth, gender, and a short bio. There\'s also a password change section where they have to enter their current password first before setting a new one, as an extra security measure.', True)

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, '3.3.5 Administrative Panel', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_body_text(doc, 'The admin panel is arguably the biggest part of the whole project. It\'s where the gym owner or manager controls everything about the website. We gave it its own layout with a sidebar navigation on the left and a content area on the right.', True)

add_body_text(doc, 'The admin dashboard is the first thing you see after logging in. It shows four main stat cards at the top: total users, active exercises, trainers, and unread messages. Each card has a growth indicator and a little icon. Below that, there\'s a section showing the most recent contact messages and newly registered users so the admin can quickly see what\'s new.', True)

add_body_text(doc, 'For managing exercises, the admin can add a new exercise through a pretty big form. It asks for the name in English and Kurdish, a short description and full description in both languages, a YouTube video URL for demonstration, difficulty level (dropdown with beginner/intermediate/advanced), which muscle group it targets, what equipment is needed, duration in minutes, calories burned, detailed instructions in both languages, tips in both languages, an image upload, and checkboxes for featured and beginner-friendly flags. Editing and deleting work the same way but with the form pre-filled with existing data.', True)

add_body_text(doc, 'The trainers management page is similar - you can add trainers with all their info including a bio in both languages, their specialization, years of experience, certifications, phone and email, avatar and cover image uploads, and social media links for Instagram, Facebook, and YouTube.', True)

add_body_text(doc, 'Membership plans are managed with fields for name, description in both languages, price, duration in days, and a text area for features where you enter one feature per line. The system stores these as a JSON array in the database. There\'s also a "popular" checkbox that adds a highlight to that plan on the public page.', True)

add_body_text(doc, 'The settings page is organized into tabs: General (site name, description, hero text), Contact (email, phone, address), Social (social media links), and Colors. The colors section has over 45 options that let the admin change pretty much any color on the site - primary color, background, text colors, button colors, navbar, footer, you name it. This was done through a dynamic CSS system that outputs CSS variables based on the database values.', True)

add_body_text(doc, 'The messages page shows all contact form submissions in a table with the sender\'s name, email, subject, and whether it\'s been read. Clicking on a message shows the full details and automatically marks it as read.', True)

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, '3.3.6 Bilingual Support System', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_body_text(doc, 'Making the site work in both Kurdish and English was probably the most challenging part of the whole project, mainly because of the RTL layout. When you switch to Kurdish, everything needs to flip - text aligns to the right, the navigation reads from right to left, and things like margins and paddings that were on the left need to move to the right.', True)

add_body_text(doc, 'We handled language switching through a simple system: there\'s a toggle in the navbar that lets users switch between EN and KU. When they click it, the language preference gets saved in the session (stored as "ku" or "en") and the page reloads. The HTML tag\'s dir attribute gets set to "rtl" or "ltr" accordingly, and CSS handles the rest.', True)

add_body_text(doc, 'For text that\'s hardcoded in the PHP files (like button labels and headings), we created a helper function called __() (double underscore). You call it like __("Join Now", "ئێستا بەشداربە") and it returns the right version based on the current language. It\'s simple but it works.', True)

add_body_text(doc, 'For content stored in the database, like exercise names or trainer bios, we use duplicate columns. So the games table has both a "name" column and a "name_ku" column. The getLocalized() function figures out which one to return based on the active language. If the Kurdish version is empty, it falls back to English so there\'s always something to display.', True)

add_body_text(doc, 'For fonts, we had to use Google\'s Noto Sans Arabic font alongside the regular Inter and Oswald fonts, because Kurdish script needs proper Arabic-script font support to look right.', True)

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, '3.3.7 Security Implementation', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_body_text(doc, 'Security was something we tried to think about from the start rather than adding it as an afterthought. Here\'s what we did:', True)

security_features = [
    'CSRF tokens on every form - we generate a random token, store it in the session, embed it as a hidden field, and verify it when the form comes back. If it doesn\'t match, the form gets rejected.',
    'Passwords are never stored in plain text. We use PHP\'s password_hash() with BCRYPT and a cost factor of 12. Even if someone got access to the database, they\'d just see hashed strings that would take years to crack.',
    'All database queries use prepared statements through PDO. This means user input is never directly inserted into SQL strings, which completely prevents SQL injection. We never concatenate user input into queries.',
    'Output escaping everywhere - any time we display user-supplied data on a page, it goes through our e() function which calls htmlspecialchars(). This prevents XSS attacks where someone might try to inject malicious JavaScript through a form field.',
    'File uploads are restricted to image types only (jpg, jpeg, png, gif, webp) and limited to 5MB. Uploaded files get renamed to random strings so attackers can\'t predict file paths or overwrite existing files.',
    'User and admin sessions are completely separate. Even if somehow a regular user session got elevated, they couldn\'t access admin functions because the code checks for admin-specific session variables.',
    'We validate inputs on the server side even though there\'s also client-side validation in JavaScript. Client-side validation can always be bypassed, so the server checks are what really matter.',
]

for feature in security_features:
    add_bullet_point(doc, feature)

add_empty_lines(doc, 1)

add_formatted_paragraph(doc, '3.3.8 Responsive Design', font_size=16, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=8)

add_body_text(doc, 'Making the site work well on different screen sizes was something we tested throughout development, not just at the end. We\'d regularly pull up Chrome\'s device toolbar and check how things looked on phone and tablet sizes.', True)

add_body_text(doc, 'Bootstrap\'s grid system did most of the heavy lifting here. We used classes like col-lg-4 col-md-6 to make cards go from three columns on desktop to two on tablet to one on mobile. The navigation collapses into a hamburger menu on smaller screens, which is standard Bootstrap behavior but we customized the animation a bit.', True)

add_body_text(doc, 'We also used CSS custom properties (variables) for a lot of the styling. This made it easy to keep colors and spacing consistent, and it\'s also what powers the admin color customization feature since changing a CSS variable instantly affects every element that uses it.', True)


# 3.4 Project Publishing
add_section_heading(doc, 'Project Publishing', '3.4')

add_body_text(doc, 'Right now FitZone runs locally through XAMPP, which works perfectly for development and demonstration. But it\'s designed to be deployed to a real web server too. Any hosting that supports PHP 8.0 or higher and MySQL 5.7 or higher would work fine.', True)

add_body_text(doc, 'To deploy it, you\'d basically need to upload all the project files to the hosting via FTP or a file manager, create a new database and import our SQL setup files through phpMyAdmin (which most hosts provide), and then update the config.php file with the production database credentials and the actual domain URL. It\'s not a complicated process if you\'re familiar with web hosting.', True)

add_body_text(doc, 'We think FitZone would work well for any small to medium sized gym that wants an online presence. The bilingual feature makes it especially useful for gyms in the Kurdistan Region since their members can use the site in their own language. And because it\'s all self-hosted, there are no monthly software fees - just the hosting cost, which can be as low as $3-5 per month for a basic plan.', True)


# ============================================================
# CHAPTER FOUR: CONCLUSION AND FUTURE WORK
# ============================================================
add_page_break(doc)

add_chapter_title(doc, 'CHAPTER FOUR', 'CONCLUSION AND FUTURE WORK')

# 4.1 Conclusion
add_section_heading(doc, 'Conclusion', '4.1')

add_body_text(doc, 'Looking back at what we\'ve built, we\'re honestly pretty proud of how FitZone turned out. When we started this project, we had a rough idea of what we wanted but we weren\'t sure we could pull it all off. There were definitely moments where we got frustrated - the bilingual RTL support alone gave us headaches for a solid week - but we pushed through and the end result is a system that we think could genuinely be useful for a real gym.', True)

add_body_text(doc, 'The main thing we set out to do was create a bilingual gym website that works in both Kurdish and English, and we accomplished that. Every page, every piece of content, every button and label can be viewed in either language, and the layout properly adjusts for right-to-left Kurdish text. As far as we know, there\'s no other gym management system out there that offers Kurdish language support, so we\'re happy we filled that gap.', True)

add_body_text(doc, 'Here\'s a quick summary of what we achieved:', True)

achievements = [
    'A fully working bilingual website with both English and Kurdish, including proper RTL layout that we tested extensively to make sure nothing looked broken when switching languages.',
    'An exercise database system that can hold hundreds of exercises with all the details a gym-goer would want - difficulty, muscles worked, equipment, video demonstrations, instructions, and tips.',
    'A secure authentication system that we\'re fairly confident in. Passwords are properly hashed, forms are protected against CSRF, queries are parameterized, and all output is escaped.',
    'A user dashboard where members can create their own workout playlists, keep a daily training journal, and manage their personal information.',
    'A comprehensive admin panel with management interfaces for 10 different types of content, plus a settings system with over 45 customizable color options.',
    'Responsive design that we tested on various screen sizes and it looks good on everything from a large monitor down to a phone screen.',
    'A clean codebase with reusable helper functions, a singleton database connection, and organized file structure that would make it reasonably easy for another developer to pick up and work on.',
]

for achievement in achievements:
    add_bullet_point(doc, achievement)

add_body_text(doc, 'Were there challenges? Absolutely. The biggest one was the bilingual system - it\'s easy to make a page display in two languages, but making sure forms, error messages, flash notifications, and dynamic content all work correctly in both directions was much harder than expected. Another challenge was scope management; we kept wanting to add "just one more feature" and had to learn to say no to ourselves so we could actually finish on time.', True)

add_body_text(doc, 'All in all, this project taught us way more than just PHP and MySQL. We learned how to plan a real software project, work as a team, make design decisions with trade-offs, and push through when things got difficult. These are skills that will stick with us well beyond this diploma.', True)

# 4.2 Future Works
add_section_heading(doc, 'Future Works', '4.2')

add_body_text(doc, 'There are quite a few things we would love to add to FitZone if we had more time. Some of these are features we planned from the start but couldn\'t fit into the timeline, and others are ideas that came up during development:', True)

future_works = [
    'Online payments - right now the plans page just shows prices, but it would be great if members could actually pay through the website using services like FIB or FastPay that people in Kurdistan use.',
    'A mobile app would be the natural next step. We could build one with React Native or Flutter that connects to the same database and lets users track their workouts from their phone without opening a browser.',
    'Weight progress charts - we already collect weight data in the daily notes, so it would be cool to show users a graph of how their weight has changed over time. Same for workout frequency.',
    'A notification system that sends emails for things like membership renewals coming up, inactivity reminders, or new tips posted.',
    'Video uploads for trainers - right now exercises link to YouTube, but it would be better if trainers could upload their own demonstration videos directly.',
    'A real-time chat feature so members could message trainers directly through the site for quick advice or scheduling.',
    'QR code check-in for gym attendance tracking. Members would scan a code when they arrive and the system would record their visit.',
    'Some kind of AI recommendation system that suggests exercises based on the user\'s goals, fitness level, and past workout history.',
    'Support for multiple gym branches, so a gym chain could manage all their locations from one system.',
    'Adding Arabic as a third language, which would let the system be used in other parts of Iraq too.',
]

for work in future_works:
    add_bullet_point(doc, work)


# ============================================================
# REFERENCES
# ============================================================
add_page_break(doc)

add_formatted_paragraph(doc, 'REFERENCES', font_size=22, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

references = [
    'Nixon, R. (2021). Learning PHP, MySQL & JavaScript: A Step-by-Step Guide to Creating Dynamic Websites. 6th Edition. O\'Reilly Media.',
    'Duckett, J. (2014). HTML and CSS: Design and Build Websites. John Wiley & Sons.',
    'Duckett, J. (2014). JavaScript and jQuery: Interactive Front-End Web Development. John Wiley & Sons.',
    'Welling, L. and Thomson, L. (2016). PHP and MySQL Web Development. 5th Edition. Addison-Wesley Professional.',
    'Spurlock, J. (2013). Bootstrap: Responsive Web Development. O\'Reilly Media.',
    'W3Techs (2024). Usage Statistics of Server-side Programming Languages for Websites. Available at: https://w3techs.com/technologies/overview/programming_language',
    'PHP Documentation (2024). PHP: Hypertext Preprocessor. Available at: https://www.php.net/docs.php',
    'MySQL Documentation (2024). MySQL 8.0 Reference Manual. Available at: https://dev.mysql.com/doc/refman/8.0/en/',
    'Bootstrap Documentation (2024). Bootstrap 5 Documentation. Available at: https://getbootstrap.com/docs/5.3/',
    'MDN Web Docs (2024). JavaScript Reference. Mozilla Developer Network. Available at: https://developer.mozilla.org/en-US/docs/Web/JavaScript',
    'OWASP Foundation (2023). OWASP Top Ten Web Application Security Risks. Available at: https://owasp.org/www-project-top-ten/',
    'Apache Friends (2024). XAMPP - Apache + MariaDB + PHP + Perl. Available at: https://www.apachefriends.org/',
]

for i, ref in enumerate(references):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

    run = p.add_run(f'[{i+1}] {ref}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'FitZone_Research_Project_Raparin_Institute.docx')
doc.save(output_path)
print(f'Document saved successfully to: {output_path}')
print('Done!')
