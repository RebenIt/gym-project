#!/usr/bin/env python3
"""Insert User Flow Diagram in section 3.2 (Project Design)."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document('Final_Research_Project_review.docx')
body = doc.element.body

# Find the position after "3.2.3 Database Design" content (after the ER diagram caption)
# We need to find Figure 3.2 caption and the end of 3.2.3 section, then insert before 3.3 Project Implementation

# Find "3.3 Project Implementation" heading
target_elem = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text == '3.3 Project Implementation':
        target_elem = p._element
        print(f"Found target: 3.3 Project Implementation at paragraph {i}")
        break

if target_elem is None:
    print("ERROR: Could not find 3.3 Project Implementation")
    exit(1)

# Build new content elements
new_paragraphs = []

# 1. Add heading "3.2.4 User Flow and Login Coverage"
h = doc.add_heading('3.2.4 User Flow and Login Coverage', level=3)
new_paragraphs.append(h._element)

# 2. Add description paragraph
desc = doc.add_paragraph()
run = desc.add_run("This section illustrates the complete user journey through the FitZone website, from the moment a visitor lands on the homepage to the post-login experience. The diagram below shows what each type of user (visitor, member, and admin) can access, and how the login process determines what features become available.")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
desc.paragraph_format.line_spacing = 1.5
desc.paragraph_format.space_after = Pt(6)
new_paragraphs.append(desc._element)

desc2 = doc.add_paragraph()
run = desc2.add_run("Visitors can browse all public pages including the homepage, exercises, trainers, membership plans, fitness tips, and contact form without needing an account. They can also switch between English and Kurdish languages at any time. When a user logs in as a regular member, they gain access to the user dashboard with personal features like workout lists, daily training notes, and profile management. When an admin logs in, they reach the admin control panel where they can manage every aspect of the website including exercises, trainers, plans, tips, certificates, services, contact messages, users, settings, colors, and pages.")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
desc2.paragraph_format.line_spacing = 1.5
desc2.paragraph_format.space_after = Pt(12)
new_paragraphs.append(desc2._element)

# 3. Add the diagram image
img_p = doc.add_paragraph()
img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = img_p.add_run()
run.add_picture('assets/images/user_flow_diagram.png', width=Inches(6.0))
new_paragraphs.append(img_p._element)

# 4. Add caption
cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cap.add_run('Figure 3.14: User Flow and Login Coverage Diagram')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.bold = True
cap.paragraph_format.space_after = Pt(16)
new_paragraphs.append(cap._element)

# Move all new paragraphs to position before "3.3 Project Implementation"
for elem in reversed(new_paragraphs):
    body.remove(elem)
    target_elem.addprevious(elem)

print(f"Inserted {len(new_paragraphs)} elements before '3.3 Project Implementation'")

# Now update the List of Figures - add Figure 3.14 entry
# Find LIST OF FIGURES heading and the last entry, insert new one
list_heading = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'LIST OF FIGURES' and 'Heading' in p.style.name:
        list_heading = p._element
        break

# Find Figure 3.13 entry and insert Figure 3.14 after it
fig_13_elem = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('Figure 3.13:'):
        fig_13_elem = p._element

if fig_13_elem is not None:
    from docx.shared import Inches as In
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

    # Create new entry: Figure 3.14: User Flow and Login Coverage Diagram
    new_p = doc.add_paragraph()
    tab_stops = new_p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(In(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    run1 = new_p.add_run('Figure 3.14: ')
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run1.bold = True

    run2 = new_p.add_run('User Flow and Login Coverage Diagram')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    run3 = new_p.add_run('\t18')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)

    new_p.paragraph_format.space_after = Pt(6)
    new_p.paragraph_format.line_spacing = 1.5

    # Move to position
    new_elem = new_p._element
    body.remove(new_elem)
    fig_13_elem.addnext(new_elem)
    print("Added Figure 3.14 entry to List of Figures")

doc.save('Final_Research_Project_review.docx')
print("\nDone! User Flow Diagram added to section 3.2.4")
