#!/usr/bin/env python3
"""
Batch update Word document:
1. Insert new screenshots (Login, Register update)
2. Remove unwanted placeholders (Workout Lists, Daily Notes, Admin Colors)
"""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('FitZone_Research_Project_Raparin_Institute_v2.docx')

# === STEP 1: Remove unwanted placeholders ===
remove_captions = [
    'User Workout Lists',
    'User Daily Notes',
    'Admin Color Management',
]

tables_to_remove = []
for caption_search in remove_captions:
    for ti, table in enumerate(doc.tables):
        try:
            cell_text = table.cell(0, 0).text
        except:
            continue
        if caption_search in cell_text:
            tbl_element = table._element
            parent = tbl_element.getparent()
            tbl_index = list(parent).index(tbl_element)

            # Remove table
            parent.remove(tbl_element)

            # Remove caption paragraph after it
            elements = list(parent)
            if tbl_index < len(elements):
                next_elem = elements[tbl_index]
                if next_elem.tag.endswith('}p'):
                    all_text = ''.join(node.text or '' for node in next_elem.iter())
                    if 'Figure' in all_text:
                        parent.remove(next_elem)

            print(f"Removed placeholder: {caption_search}")
            break

# === STEP 2: Insert new screenshots ===
insertions = {
    'User Login Page': (
        'assets/images/screenshots/Screenshot (751).png',
        'Figure 3.6: User Login Page'
    ),
}

for caption_search, (img_path, new_caption) in insertions.items():
    for ti, table in enumerate(doc.tables):
        try:
            cell_text = table.cell(0, 0).text
        except:
            continue
        if caption_search in cell_text:
            tbl_element = table._element
            parent = tbl_element.getparent()
            tbl_index = list(parent).index(tbl_element)

            # Remove placeholder table
            parent.remove(tbl_element)

            # Remove old caption
            elements = list(parent)
            if tbl_index < len(elements):
                next_elem = elements[tbl_index]
                if next_elem.tag.endswith('}p'):
                    all_text = ''.join(node.text or '' for node in next_elem.iter())
                    if 'Figure' in all_text:
                        parent.remove(next_elem)

            # Add image
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(5.5))
            img_elem = p._element

            # Add caption
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(new_caption)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            cap.paragraph_format.space_after = Pt(16)
            cap_elem = cap._element

            # Move to correct position
            body = doc.element.body
            body.remove(cap_elem)
            body.remove(img_elem)
            elements_list = list(body)
            if tbl_index < len(elements_list):
                elements_list[tbl_index].addprevious(img_elem)
                img_idx = list(body).index(img_elem)
                list(body)[img_idx].addnext(cap_elem)

            print(f"Inserted: {new_caption}")
            break

# === STEP 3: Update Register screenshot with newer version ===
# Find the existing Figure 3.5 image and its caption, replace with new screenshot
# We'll find the caption paragraph and the image paragraph before it
for i, p in enumerate(doc.paragraphs):
    if 'Figure 3.5: User Registration Page' in p.text:
        # The image paragraph should be just before this caption
        body = doc.element.body
        cap_elem = p._element
        cap_idx = list(body).index(cap_elem)

        # Previous element should be the image
        prev_elem = list(body)[cap_idx - 1]

        # Remove old image paragraph
        body.remove(prev_elem)

        # Add new image
        new_p = doc.add_paragraph()
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = new_p.add_run()
        run.add_picture('assets/images/screenshots/Screenshot (752).png', width=Inches(5.5))
        new_elem = new_p._element

        # Move before caption
        body.remove(new_elem)
        cap_elem.addprevious(new_elem)

        print("Updated: Figure 3.5 Register Page with newer screenshot")
        break

doc.save('FitZone_Research_Project_Raparin_Institute_v2.docx')
print("\nAll updates complete!")
