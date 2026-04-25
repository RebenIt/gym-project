#!/usr/bin/env python3
"""Insert ER diagram into the Word document."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('FitZone_Research_Project_Raparin_Institute_v2.docx')

for ti, table in enumerate(doc.tables):
    try:
        cell_text = table.cell(0, 0).text
    except:
        continue
    if 'Database Entity-Relationship Diagram' in cell_text:
        tbl_element = table._element
        parent = tbl_element.getparent()
        tbl_index = list(parent).index(tbl_element)

        parent.remove(tbl_element)

        # Remove old caption
        elements = list(parent)
        if tbl_index < len(elements):
            next_elem = elements[tbl_index]
            if next_elem.tag.endswith('}p'):
                all_text = ''.join(node.text or '' for node in next_elem.iter())
                if 'Figure' in all_text:
                    parent.remove(next_elem)

        # Add image
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture('assets/images/er_diagram.png', width=Inches(6.0))
        img_elem = p._element

        # Add caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run('Figure 3.2: Database Entity-Relationship Diagram')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True
        cap.paragraph_format.space_after = Pt(16)
        cap_elem = cap._element

        body = doc.element.body
        body.remove(cap_elem)
        body.remove(img_elem)
        elements_list = list(body)
        if tbl_index < len(elements_list):
            elements_list[tbl_index].addprevious(img_elem)
            img_idx = list(body).index(img_elem)
            list(body)[img_idx].addnext(cap_elem)

        print("Inserted: ER Diagram")
        break

doc.save('FitZone_Research_Project_Raparin_Institute_v2.docx')

# Final check
doc2 = Document('FitZone_Research_Project_Raparin_Institute_v2.docx')
img_count = sum(1 for rel in doc2.part.rels.values() if 'image' in rel.reltype)
placeholder_count = sum(1 for t in doc2.tables
                       for _ in [1] if 'Screenshot Placeholder' in (t.cell(0,0).text if t.rows else ''))
print(f"\nTotal images in document: {img_count}")
print(f"Remaining placeholders: {placeholder_count}")
print("Done!")
