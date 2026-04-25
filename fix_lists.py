#!/usr/bin/env python3
"""
Fix List of Figures and List of Tables to match template format:
Figure 3.1: Title.......................... 7
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER

doc = Document('Final_Research_Project_review.docx')

# All figures with their estimated page numbers (sequential order in doc)
figures = [
    ('Figure 3.1', 'System Architecture Diagram', 15),
    ('Figure 3.2', 'Database Entity-Relationship Diagram', 16),
    ('Figure 3.3', 'Website Flowchart', 17),
    ('Figure 3.4a', 'Home Page - Hero Section, Stats & Services', 18),
    ('Figure 3.4b', 'Home Page - Membership Plans', 19),
    ('Figure 3.4c', 'Home Page - Trainers Section', 20),
    ('Figure 3.4d', 'Home Page - Reviews & Footer', 21),
    ('Figure 3.5', 'User Registration Page', 22),
    ('Figure 3.6', 'User Login Page', 23),
    ('Figure 3.7', 'Exercises Page with Filters', 24),
    ('Figure 3.8', 'Trainers Page', 25),
    ('Figure 3.9', 'Membership Plans Page', 26),
    ('Figure 3.10', 'User Dashboard', 27),
    ('Figure 3.11', 'Admin Dashboard', 28),
    ('Figure 3.12', 'Admin Games Management', 29),
    ('Figure 3.13', 'Admin Settings Page', 30),
    ('Figure 3.15', 'Contact Page', 31),
    ('Figure 3.16', 'Tips / Blog Page', 32),
    ('Figure 3.17', 'Admin Trainers Management', 33),
    ('Figure 3.18', 'Admin Plans Management', 34),
    ('Figure 3.19', 'Admin Messages Page', 35),
    ('Figure 3.22', 'User Profile Page', 36),
]

# All tables with their estimated page numbers
tables_list = [
    ('Table 2.1', 'Software Requirements', 11),
    ('Table 2.2', 'Hardware Requirements', 12),
    ('Table 3.1', 'Project Plan and Timeline', 13),
    ('Table 3.2', 'Database Tables Overview', 14),
    ('Table 3.3', 'Users Table Structure', 14),
    ('Table 3.4', 'Acronyms List', 8),
    ('Table 3.5', 'Figure Index', 9),
    ('Table 3.6', 'Table Index', 10),
]

# Find and remove the figure list table (Table 2) and table list table (Table 3)
tables_to_remove = []
for ti, table in enumerate(doc.tables):
    if not table.rows:
        continue
    first_cell = table.cell(0, 0).text.strip()
    if first_cell == 'Index':
        # Check second cell
        second_cell = table.cell(0, 1).text.strip() if len(table.rows[0].cells) > 1 else ''
        if 'Figure' in second_cell or 'Table' in second_cell:
            tables_to_remove.append((ti, second_cell, table))
            print(f"Will remove Table {ti}: {first_cell} | {second_cell}")

# Helper to add a list entry with tab leader
def add_list_entry(insert_after_elem, label, title, page_num, body):
    p = doc.add_paragraph()
    # Set up tab stop with dot leader at right margin
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    # Add label (bold)
    run1 = p.add_run(f'{label}: ')
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run1.bold = True

    # Add title
    run2 = p.add_run(title)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    # Tab and page number
    run3 = p.add_run(f'\t{page_num}')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)

    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5

    # Move to position
    new_elem = p._element
    body.remove(new_elem)
    insert_after_elem.addnext(new_elem)
    return new_elem

# Now process each table to remove
body = doc.element.body

# Process figures table
for ti, label, table in tables_to_remove:
    is_figure_list = 'Figure' in label
    tbl_element = table._element
    parent = tbl_element.getparent()
    tbl_index = list(parent).index(tbl_element)

    # Remove the table
    parent.remove(tbl_element)

    # Find the LIST OF FIGURES or LIST OF TABLES heading right before this position
    # We'll insert new paragraphs after that heading
    heading_elem = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if is_figure_list and text == 'LIST OF FIGURES':
            heading_elem = p._element
            break
        elif not is_figure_list and text == 'LIST OF TABLES':
            heading_elem = p._element
            break

    if heading_elem is None:
        print(f"WARNING: Heading not found for {label}")
        continue

    items = figures if is_figure_list else tables_list
    insert_after = heading_elem

    for fig_label, fig_title, page in items:
        insert_after = add_list_entry(insert_after, fig_label, fig_title, page, body)

    print(f"Added {len(items)} entries for {label}")

doc.save('Final_Research_Project_review.docx')
print("\nDone!")
