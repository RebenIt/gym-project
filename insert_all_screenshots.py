#!/usr/bin/env python3
"""Insert all remaining screenshots into the Word document."""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('FitZone_Research_Project_Raparin_Institute_v2.docx')
sdir = 'assets/images/screenshots/'

# Map: placeholder caption search text -> (image path, new caption)
insertions = {
    'Exercises Page with Filters': (sdir + 'Screenshot (753).png', 'Figure 3.7: Exercises Page with Filters'),
    'Trainers Page': (sdir + 'Screenshot (754).png', 'Figure 3.8: Trainers Page'),
    'Tips/Blog Page': (sdir + 'Screenshot (755).png', 'Figure 3.16: Tips / Blog Page'),
    'Contact Page': (sdir + 'Screenshot (756).png', 'Figure 3.15: Contact Page'),
    'Membership Plans Page': (sdir + 'Screenshot (760).png', 'Figure 3.9: Membership Plans Page'),
    'User Dashboard': (sdir + 'Screenshot (761).png', 'Figure 3.10: User Dashboard'),
    'User Profile Page': (sdir + 'Screenshot (763).png', 'Figure 3.22: User Profile Page'),
    'Admin Dashboard': (sdir + 'Screenshot (764).png', 'Figure 3.11: Admin Dashboard'),
    'Admin Games Management': (sdir + 'Screenshot (765).png', 'Figure 3.12: Admin Games Management'),
    'Admin Trainers Management': (sdir + 'Screenshot (767).png', 'Figure 3.17: Admin Trainers Management'),
    'Admin Plans Management': (sdir + 'Screenshot (768).png', 'Figure 3.18: Admin Plans Management'),
    'Admin Settings Page': (sdir + 'Screenshot (769).png', 'Figure 3.13: Admin Settings Page'),
    'Admin Messages Page': (sdir + 'Screenshot (770).png', 'Figure 3.19: Admin Messages Page'),
}

# Also remove these placeholders (no screenshot needed)
remove_only = [
    'Kurdish Language View',
    'Mobile Responsive View',
]

inserted_count = 0
removed_count = 0

# Process removals first
for caption_search in remove_only:
    for ti, table in enumerate(doc.tables):
        try:
            cell_text = table.cell(0, 0).text
        except:
            continue
        if caption_search in cell_text:
            tbl_element = table._element
            parent = tbl_element.getparent()
            tbl_index = list(parent).index(tbl_element)
            parent.remove(tbl_element)
            # Remove caption
            elements = list(parent)
            if tbl_index < len(elements):
                next_elem = elements[tbl_index]
                if next_elem.tag.endswith('}p'):
                    all_text = ''.join(node.text or '' for node in next_elem.iter())
                    if 'Figure' in all_text:
                        parent.remove(next_elem)
            removed_count += 1
            print(f"Removed: {caption_search}")
            break

# Process insertions
for caption_search, (img_path, new_caption) in insertions.items():
    found = False
    for ti, table in enumerate(doc.tables):
        try:
            cell_text = table.cell(0, 0).text
        except:
            continue
        if caption_search in cell_text:
            found = True
            tbl_element = table._element
            parent = tbl_element.getparent()
            tbl_index = list(parent).index(tbl_element)

            # Remove placeholder table
            parent.remove(tbl_element)

            # Remove old caption
            elements = list(parent)
            if tbl_index < len(elements):
                next_elem = elements[tbl_index]
                if next_elem.tag.endswith('}p'):
                    all_text = ''.join(node.text or '' for node in next_elem.iter())
                    if 'Figure' in all_text:
                        parent.remove(next_elem)

            # Add image
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(5.5))
            img_elem = p._element

            # Add caption
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(new_caption)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            cap.paragraph_format.space_after = Pt(16)
            cap_elem = cap._element

            # Move to correct position
            body = doc.element.body
            body.remove(cap_elem)
            body.remove(img_elem)
            elements_list = list(body)
            if tbl_index < len(elements_list):
                elements_list[tbl_index].addprevious(img_elem)
                img_idx = list(body).index(img_elem)
                list(body)[img_idx].addnext(cap_elem)

            inserted_count += 1
            print(f"Inserted: {new_caption}")
            break

    if not found:
        print(f"WARNING: Placeholder not found for '{caption_search}'")

doc.save('FitZone_Research_Project_Raparin_Institute_v2.docx')
print(f"\nDone! Inserted: {inserted_count}, Removed: {removed_count}")

# Final count
doc2 = Document('FitZone_Research_Project_Raparin_Institute_v2.docx')
img_count = sum(1 for rel in doc2.part.rels.values() if 'image' in rel.reltype)
placeholder_count = sum(1 for t in doc2.tables
                       for _ in [1] if 'Screenshot Placeholder' in (t.cell(0,0).text if t.rows else ''))
print(f"Total images: {img_count}")
print(f"Remaining placeholders: {placeholder_count}")
