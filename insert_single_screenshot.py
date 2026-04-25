#!/usr/bin/env python3
"""Insert a single screenshot into the Word document, replacing a placeholder."""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('FitZone_Research_Project_Raparin_Institute_v2.docx')

# Map: placeholder search text -> (image file, new caption)
replacement = {
    'User Registration Page': (
        'assets/images/screenshots/Screenshot (750).png',
        'Figure 3.5: User Registration Page'
    ),
}

for caption_search, (img_path, new_caption) in replacement.items():
    for ti, table in enumerate(doc.tables):
        try:
            cell_text = table.cell(0, 0).text
        except:
            continue

        if caption_search in cell_text:
            print(f"Found placeholder: '{caption_search}'")
            tbl_element = table._element
            parent = tbl_element.getparent()
            tbl_index = list(parent).index(tbl_element)

            # Remove placeholder table
            parent.remove(tbl_element)

            # Remove old caption paragraph
            elements = list(parent)
            if tbl_index < len(elements):
                next_elem = elements[tbl_index]
                if next_elem.tag.endswith('}p'):
                    all_text = ''.join(node.text or '' for node in next_elem.iter())
                    if 'Figure' in all_text and caption_search.split()[0] in all_text:
                        parent.remove(next_elem)

            # Add image
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(5.5))
            img_elem = p._element

            # Add caption
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(new_caption)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            cap.paragraph_format.space_after = Pt(16)
            cap_elem = cap._element

            # Move to correct position
            body = doc.element.body
            body.remove(cap_elem)
            body.remove(img_elem)
            elements_list = list(body)
            if tbl_index < len(elements_list):
                elements_list[tbl_index].addprevious(img_elem)
                img_idx = list(body).index(img_elem)
                list(body)[img_idx].addnext(cap_elem)

            print(f"  Inserted: {new_caption}")
            break

doc.save('FitZone_Research_Project_Raparin_Institute_v2.docx')
print("Done!")
