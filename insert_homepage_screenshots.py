#!/usr/bin/env python3
"""Insert homepage screenshots into the Word document at Figure 3.4 placeholder."""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('FitZone_Research_Project_Raparin_Institute_v2.docx')

screenshots_dir = 'assets/images/screenshots/'

# Map: placeholder caption text -> list of images to insert
# We'll find tables that contain placeholder text and replace them
replacements = {
    'Home Page Design': [
        (screenshots_dir + 'homepage_hero.png', 'Figure 3.4a: Home Page - Hero Section, Stats & Services'),
        (screenshots_dir + 'homepage_plans.png', 'Figure 3.4b: Home Page - Membership Plans'),
        (screenshots_dir + 'homepage_trainers.png', 'Figure 3.4c: Home Page - Trainers Section'),
        (screenshots_dir + 'homepage_footer.png', 'Figure 3.4d: Home Page - Reviews & Footer'),
    ],
}

tables_to_remove = []
insert_positions = {}

# Find placeholder tables and their positions
for ti, table in enumerate(doc.tables):
    try:
        cell_text = table.cell(0, 0).text
    except:
        continue

    for caption, images in replacements.items():
        if caption in cell_text:
            print(f"Found placeholder: '{caption}' in table {ti}")
            # Get the table's parent element and position
            tbl_element = table._element
            parent = tbl_element.getparent()

            # Find the caption paragraph after this table
            # We need to insert images before the caption paragraph
            tbl_index = list(parent).index(tbl_element)

            # Remove the placeholder table
            parent.remove(tbl_element)

            # Also find and remove the caption paragraph (Figure 3.4: Home Page Design)
            # It should be the next element after the table was
            elements = list(parent)
            if tbl_index < len(elements):
                next_elem = elements[tbl_index]
                if next_elem.tag.endswith('}p'):
                    text = next_elem.text or ''
                    # Get all text from the paragraph
                    all_text = ''.join(node.text or '' for node in next_elem.iter())
                    if 'Figure 3.4' in all_text:
                        parent.remove(next_elem)
                        print(f"  Removed old caption paragraph")

            # Now insert images at the position
            from docx.oxml.ns import qn, nsdecls
            from docx.oxml import parse_xml
            from lxml import etree

            insert_idx = tbl_index
            for img_path, img_caption in images:
                # Create a new paragraph for the image
                new_para = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
                parent.insert(insert_idx, new_para)
                insert_idx += 1

                # We need to use python-docx's API to add the image
                # So we'll add paragraphs at the end and move them

            # Alternative approach: add images using python-docx API then move elements
            # Let's use a simpler approach - add at end then reorder

            added_elements = []
            for img_path, img_caption in images:
                # Add image paragraph
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5.8))
                added_elements.append(p._element)

                # Add caption paragraph
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(img_caption)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True
                cap.paragraph_format.space_after = Pt(16)
                added_elements.append(cap._element)

                print(f"  Added: {img_caption}")

            # Move added elements to the correct position
            body = doc.element.body
            for elem in reversed(added_elements):
                body.remove(elem)
                # Insert at the saved position
                elements_list = list(body)
                if tbl_index < len(elements_list):
                    elements_list[tbl_index].addprevious(elem)
                else:
                    body.append(elem)

            print(f"  Moved {len(added_elements)} elements to position {tbl_index}")

doc.save('FitZone_Research_Project_Raparin_Institute_v2.docx')
print("\nDocument updated with homepage screenshots!")
